import re
import docx
import os
import pandas as pd


############################################
#          Class MyDoc
############################################
class MyDoc:
	def __init__(self):
		self.id = -1
		self.type = "UNKNOWN"
		self.results_col = -1
		self.table_num = -1
		self.doc = -1
		self.sample = -1
		self.producer = -1
		self.client = -1

############################################
#          GET FORMAT
############################################
def get_format(doc):
	mydoc = MyDoc()
	mydoc.doc = doc

	paragraphs = enumerate(doc.paragraphs)
	for i, paragraph in paragraphs:
		if (re.search("ПРОТОКОЛ", paragraph.text, re.IGNORECASE) and mydoc.type == "UNKNOWN"):
			index = paragraph.text.find("№")
			if (index != -1):
				mydoc.id = paragraph.text[index+1 : -2].strip()
			mydoc.type = "ПРОТОКОЛ"

		elif (re.search("ПАСПОРТ", paragraph.text, re.IGNORECASE) and mydoc.type == "UNKNOWN"):
			index = paragraph.text.find("№")
			if (index != -1):
				mydoc.id = paragraph.text[index+1 : -2].strip()		
			mydoc.type = "ПАСПОРТ"
			mydoc.table_num = 0
		

		# Get sample, client and producer info
		if (re.search("найменування", paragraph.text, re.IGNORECASE) and mydoc.sample == -1):
			index = paragraph.text.find(":")
			if (index != -1):
				mydoc.sample = paragraph.text[index+1 : -1]
				mydoc.sample = clear(mydoc.sample)
				
		if (re.search("Виробник", paragraph.text, re.IGNORECASE) and mydoc.producer == -1):
			index = paragraph.text.find(":")
			if (index != -1):
				mydoc.producer = paragraph.text[index+1 : -1]
				mydoc.producer = clear(mydoc.producer)
				
		if (re.search("Замовник", paragraph.text, re.IGNORECASE) and mydoc.client == -1):
			index = paragraph.text.find(":")
			if (index != -1):
				mydoc.client = paragraph.text[index+1 : -1]
				mydoc.client = clear(mydoc.client)

		# Get results table for "ПРОТОКОЛ"
		if(re.search("Результати випроб", paragraph.text, re.IGNORECASE)):
			if(re.search("таблиц", paragraph.text, re.IGNORECASE)):
				mydoc.table_num = int(paragraph.text[-2])
				break

			else:
				for k in range(1,5):
					new_par = doc.paragraphs[i+k]
					if(re.search("таблиц", new_par.text, re.IGNORECASE)):
						mydoc.table_num = int(new_par.text[-1])
						break
					else: 
						mydoc.table_num = 2
						break 		
	
	# Get number of results column
	tab_titles = doc.tables[mydoc.table_num].rows[0].cells
	for k in range(0,len(tab_titles),1):
		if(re.search("результат", tab_titles[k].text, re.IGNORECASE)):
			mydoc.results_col = k
			break
		else:
			mydoc.results_col = -1
	
	
	return (mydoc)

############################################
#   Remove new line at the end of string
############################################
def clear(string):
	string = string.strip()
	string = string.replace('\r',';')
	string = string.replace('\t','')
	string = string.replace('\n','')
	return(string)
				
############################################
#       GET NAMES OF FILES IN DIRECTORY
############################################
def get_names(path):

	files_list = []
	import os
	i = 0
	for root, dirs, files in os.walk(os.path.abspath(path)):
		for file in files:

			files_list.append(os.path.join(root, file))
			i = i + 1
	return files_list

############################################
#       FIND PROPERTY IN PROTOCOL
############################################
'''
def find_property(substr, mydoc):
	table = mydoc.doc.tables[mydoc.table_num]

	for row in table.rows:
		prop_name = row.cells[1].text
		if (re.search(substr, prop_name, re.IGNORECASE)):
			result = row.cells[mydoc.results_col].text
			result = result.strip()
			result = result.replace('\n', '')
			result = result.replace('\t', '')
			return(result)
	
	return("na")
'''
def find_properties(df,prop_dict,mydoc,i):
	table = mydoc.doc.tables[mydoc.table_num]
	
	for row in table.rows:
		prop_name = row.cells[1].text
		for key, prop in prop_dict.items():
			if (re.search(prop, prop_name, re.IGNORECASE)):
				result = row.cells[mydoc.results_col].text
				result = result.strip()
				result = result.replace('\n', '')
				result = result.replace('\t', '')
				result = result.replace('\r', '')
				df.at[i,key] = result
				break
	return(df)
############################################
#          Clean chrom data
############################################
def clean_chrom_data(df):
	condition1 = df.ID.str.contains("\d{4}", regex = True, na = False)
	condition2 = df.ID.str.contains("\\\\\d{3}", regex = True, na = False)

	df1 = df[condition1]
	df2 = df[condition2]

	result = pd.concat([df1,df2], ignore_index = True)
	removed = df[~condition1]
	removed = removed[~condition2]

	removed.to_csv("chrom_removed_items.csv")
	return(result)

